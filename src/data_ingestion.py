"""Data ingestion pipeline for processing various document formats."""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any
import pandas as pd
from PIL import Image
import pytesseract
from docx import Document
import PyPDF2
from unstructured.partition.auto import partition
from unstructured.partition.xlsx import partition_xlsx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config.settings import settings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of various document formats."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
        )
    
    def process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using pdfplumber (surgical) with unstructured fallback."""
        try:
            import pdfplumber
            import re
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Layout=True helps preserve columns in resumes
                    page_text = page.extract_text(layout=True)
                    if page_text:
                        # Clean up common PDF artifacts like (cid:131)
                        page_text = re.sub(r'\(cid:\d+\)', ' ', page_text)
                        text_parts.append(page_text)
            
            if text_parts:
                logger.info(f"Successfully extracted clean text from PDF: {file_path.name}")
                return "\n".join(text_parts)
            
            # Fallback to unstructured if pdfplumber finds nothing
            logger.info(f"pdfplumber found no text in {file_path.name}, falling back to unstructured...")
            from unstructured.partition.pdf import partition_pdf
            elements = partition_pdf(filename=str(file_path), strategy="hi_res")
            return "\n".join([str(el) for el in elements])
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            return ""
    
    def process_docx(self, file_path: Path) -> str:
        """Extract text and surgically extract tables from DOCX files."""
        try:
            from docx import Document
            doc = Document(file_path)
            full_text = []
            
            # 1. Extract regular paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # 2. Surgical Table Extraction & Transposition
            for table in doc.tables:
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                
                if not data:
                    continue
                    
                df = pd.DataFrame(data)
                # Clean up empty rows/cols
                df = df.dropna(how='all').dropna(axis=1, how='all')
                
                # Atomic Transposition: ONLY if it looks like a financial/time-series table
                if len(df.columns) > 1:
                    metric_col = 0
                    
                    # Detect if any header looks like a year/date
                    header_row = [str(x).lower() for x in df.iloc[0].values]
                    is_financial = any(kw in "".join(header_row) for kw in ["202", "201", "fy", "year", "projection", "actual"])
                    
                    if is_financial:
                        full_text.append("\n[STRUCTURED DATA - ATOMIC FINANCIAL TAGGING]\n")
                        for col_idx in range(1, len(df.columns)):
                            year_header = str(df.iloc[0, col_idx]).strip() or f"Column_{col_idx}"
                            for row_idx in range(1, len(df)):
                                metric_label = str(df.iloc[row_idx, metric_col]).strip()
                                value = str(df.iloc[row_idx, col_idx]).strip()
                                if metric_label and value:
                                    full_text.append(f"[FINANCIAL_DATA] Year: {year_header} | Metric: {metric_label} | Value: {value}")
                            full_text.append("") # Atomic boundary
                    else:
                        # General table
                        full_text.append("\n[TABLE DETECTED]\n" + df.to_markdown(index=False) + "\n")
                else:
                    full_text.append("\n[TABLE DETECTED]\n" + df.to_markdown(index=False) + "\n")
            
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Surgical DOCX parsing failed for {file_path}: {e}")
            return ""
    
    def process_csv(self, file_path: Path) -> str:
        """Convert CSV to readable text format."""
        try:
            df = pd.read_csv(file_path)
            # Convert to structured text
            text = f"Data from {file_path.name}:\n"
            text += df.to_string(index=False)
            return text
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            return ""
    
    def process_image(self, file_path: Path) -> str:
        """Extract text from images using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            return ""
    
    def process_text(self, file_path: Path) -> str:
        """Process plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return ""
    
    def process_excel(self, file_path: Path) -> str:
        """Extract structured data from Excel files with auto-transposition for wide sheets."""
        try:
            df_dict = pd.read_excel(file_path, sheet_name=None)
            text = f"Spreadsheet: {file_path.name}\n"
            for sheet, df in df_dict.items():
                if len(df.columns) > 1:
                    metric_col = df.columns[0]
                    text += f"\n[STRUCTURED DATA - ATOMIC TAGGING: {sheet}]\n"
                    for col in df.columns[1:]:
                        for idx, row in df.iterrows():
                            text += f"{col} {row[metric_col]}: {row[col]}\n"
                        text += "\n"
                else:
                    text += f"\n[TABLE DETECTED - Sheet: {sheet}]\n"
                    text += df.to_markdown(index=False)
            return text.strip()
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            return ""

    def process_file(self, file_path: Path) -> str:
        """Route file to appropriate processor based on extension."""
        extension = file_path.suffix.lower()
        
        processors = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.doc': self.process_docx,
            '.xlsx': self.process_excel,
            '.xls': self.process_excel,
            '.csv': self.process_csv,
            '.txt': self.process_text,
            '.md': self.process_text,
            '.jpg': self.process_image,
            '.jpeg': self.process_image,
            '.png': self.process_image,
        }
        
        processor = processors.get(extension)
        if processor:
            logger.info(f"Processing {file_path.name} ({extension})")
            return processor(file_path)
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return ""
    
    def _reconstruct_table(self, text: str) -> str:
        """Heuristic to detect and reconstruct flattened tables into Markdown."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) < 4: return text
        
        # Check if it looks like a financial list (e.g., headers then numbers)
        headers = []
        numbers = []
        for line in lines:
            if any(yr in line for yr in ["FY2", "FY1", "202", "201"]):
                headers.append(line)
            elif any(c.isdigit() for c in line) and len(line) < 15:
                numbers.append(line)
        
        if headers and numbers and len(numbers) >= len(headers):
            # Try to build a markdown table
            try:
                table = "| Metric | " + " | ".join(headers) + " |\n"
                table += "| --- | " + " | ".join(["---"] * len(headers)) + " |\n"
                # This is a simplification; in reality we'd group numbers by metric
                # But even a flat table-like block is better than a list
                table += "| Data | " + " | ".join(numbers[:len(headers)]) + " |\n"
                return text + "\n\n### Reconstructed Table Data:\n" + table
            except:
                return text
        return text

    def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split text into chunks with metadata and context headers."""
        if not text.strip():
            return []
        
        # Reconstruct tables if detected
        processed_text = self._reconstruct_table(text)
        
        # Add a context header
        filename = metadata.get("filename", "unknown")
        contextualized_text = f"[Document: {filename}]\n{processed_text}"
        
        chunks = self.text_splitter.split_text(contextualized_text)
        return [
            {
                "content": chunk if chunk.startswith(f"[Document:") else f"[Document: {filename}]\n{chunk}",
                "metadata": {
                    **metadata,
                    "chunk_index": i,
                    "chunk_count": len(chunks)
                }
            }
            for i, chunk in enumerate(chunks)
        ]

class DataIngestionPipeline:
    """Main pipeline for ingesting and processing documents."""
    
    def __init__(self, embedding_manager=None):
        self.processor = DocumentProcessor()
        self.embedding_manager = embedding_manager # optional, just for consistency if needed later
    
    def ingest_documents(self) -> List[Dict[str, Any]]:
        """Process all documents in the raw data directory."""
        all_chunks = []
        
        for file_path in settings.RAW_DATA_DIR.rglob("*"):
            if file_path.is_file():
                # Extract text from file
                text = self.processor.process_file(file_path)
                
                if text:
                    # Create metadata
                    metadata = {
                        "source": str(file_path),
                        "filename": file_path.name,
                        "file_type": file_path.suffix.lower(),
                        "file_size": file_path.stat().st_size,
                    }
                    
                    # Chunk the text
                    chunks = self.processor.chunk_text(text, metadata)
                    all_chunks.extend(chunks)
                    
                    logger.info(f"Processed {file_path.name}: {len(chunks)} chunks")
        
        logger.info(f"Total chunks created: {len(all_chunks)}")
        return all_chunks
    
    def save_processed_data(self, chunks: List[Dict[str, Any]]) -> Path:
        """Save processed chunks to disk."""
        output_file = settings.PROCESSED_DATA_DIR / "processed_chunks.json"
        
        import json
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(chunks, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Saved {len(chunks)} chunks to {output_file}")
        return output_file

# Example usage
if __name__ == "__main__":
    pipeline = DataIngestionPipeline()
    chunks = pipeline.ingest_documents()
    pipeline.save_processed_data(chunks)
